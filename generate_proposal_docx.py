from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def add_heading(doc: Document, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    h.runs[0].font.name = "Calibri"
    return h


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.paragraph_format.space_after = Pt(3)


def add_numbered(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph(item, style="List Number")
        p.paragraph_format.space_after = Pt(3)


def main():
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    title = doc.add_paragraph("RÌNWÁ\nWeb Application Proposal")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(22)

    meta = doc.add_paragraph(
        "Prepared by: Development Team\n"
        "Date: 3 April 2026\n"
        "Estimated Timeline: 8 weeks\n"
        "Budget (Baseline): ₦2,650,000"
    )
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(16)

    add_heading(doc, "1) Project Goal & Success Criteria", 1)
    doc.add_paragraph(
        "Goal: Create a sophisticated, culture-first hospitality brand website with a professional "
        "aesthetic that presents RÌNWÁ as a trusted curator of intentional moments, backed by an "
        "admin console for easy content and event management."
    )
    add_bullets(
        doc,
        [
            "Website launches within 8 weeks with responsive performance across devices.",
            "Admin console enables non-technical content updates for key sections.",
            "SEO and structured data are implemented for discoverability.",
            "Launch quality target: no critical bugs, plus post-launch warranty support.",
        ],
    )

    add_heading(doc, "2) Scope of Work (Baseline)", 1)
    doc.add_paragraph("Pages/Sections:")
    add_bullets(
        doc,
        [
            "Home (hero carousel, brand marquee, experience section, calls-to-action)",
            "Featured Event block",
            "Past Events Gallery",
            "Media Gallery",
            "Contact/Inquiry section and social/footer blocks",
        ],
    )

    doc.add_paragraph("Included:")
    add_bullets(
        doc,
        [
            "Responsive UI/UX system (typography, colors, buttons, cards)",
            "SEO setup: page metadata, Open Graph/Twitter cards, sitemap.xml, robots.txt",
            "Structured data: Organization + Event schema",
            "Performance optimization: media loading, caching, code-splitting",
            "Accessibility: semantic HTML, focus states, aria labels, alt text",
            "Analytics setup (Google Analytics or Plausible)",
            "Staging + production deployment",
            "Admin console with role-based access and content workflows",
        ],
    )

    doc.add_paragraph("Explicitly excluded (Baseline):")
    add_bullets(
        doc,
        [
            "Custom booking or payment engine",
            "Advanced CRM/email automation integrations",
            "Multi-language support (unless separately scoped)",
            "Long-term maintenance outside included support window",
        ],
    )

    add_heading(doc, "3) Technical Approach", 1)
    add_bullets(
        doc,
        [
            "Framework: Next.js (App Router) + TypeScript",
            "Styling/UI: Tailwind CSS + modern component architecture",
            "Animation: Framer Motion",
            "Admin: React-based dashboard with content CRUD workflows",
            "Hosting/Deploy: Vercel (staging + production)",
            "Version Control: GitHub with protected main branch",
        ],
    )

    add_heading(doc, "4) Project Plan & Timeline", 1)
    add_bullets(
        doc,
        [
            "Week 1–2 — Discovery & planning, assets, and content alignment",
            "Week 3–5 — Frontend build, responsiveness, accessibility, SEO",
            "Week 4–6 — Backend + admin console development",
            "Week 7 — Integration, QA, and performance hardening",
            "Week 8 — Launch, handover, and admin onboarding",
            "Contingency: DNS/domain propagation may take 24–48 hours",
        ],
    )

    add_heading(doc, "5) Deliverables", 1)
    add_bullets(
        doc,
        [
            "Live production website + staging URL",
            "Functional admin console",
            "Source code repository and deployment pipeline",
            "SEO + structured data implementation",
            "Documentation and admin handover guide",
            "Post-launch support window",
        ],
    )

    add_heading(doc, "6) Budget", 1)
    add_bullets(
        doc,
        [
            "Website Development: ₦950,000",
            "Admin Console: ₦750,000",
            "Backend Integration: ₦500,000",
            "Testing & QA: ₦200,000",
            "Deployment & Documentation: ₦150,000",
            "Post-Launch Support: ₦100,000",
            "Total: ₦2,650,000",
        ],
    )

    add_heading(doc, "7) Payment Terms", 1)
    add_bullets(
        doc,
        [
            "50% deposit (₦1,325,000) to commence work",
            "50% balance (₦1,325,000) due at launch",
            "Quote valid for 30 days from issue date",
        ],
    )

    add_heading(doc, "8) Assumptions & Client Responsibilities", 1)
    add_bullets(
        doc,
        [
            "Client provides final copy, logos, and media assets",
            "Single approver provides feedback within 24–48 hours",
            "Two revision rounds included in baseline scope",
            "Domain and hosting access credentials provided as needed",
        ],
    )

    add_heading(doc, "9) Risks & Mitigations", 1)
    add_bullets(
        doc,
        [
            "Content delays: use placeholders and locked review dates",
            "Domain/DNS delays: launch on staging first",
            "Scope expansion: manage through change requests and separate quotes",
        ],
    )

    add_heading(doc, "10) Handover & Maintenance (Optional)", 1)
    add_bullets(
        doc,
        [
            "1-hour training session for content and admin workflows",
            "Documentation included for updates and operations",
            "Optional monthly maintenance plan available",
        ],
    )

    add_heading(doc, "11) Acceptance", 1)
    doc.add_paragraph("Authorized Signatory (RÌNWÁ Hospitality & Experiences):")
    doc.add_paragraph("Name: ____________________")
    doc.add_paragraph("Title: ____________________")
    doc.add_paragraph("Signature/Date: ____________________")
    doc.add_paragraph("")
    doc.add_paragraph("Prepared by (Development Team):")
    doc.add_paragraph("Signature/Date: ____________________")

    output = "/Users/mac/Desktop/the-badest-event/PROPOSAL_ALIGNED.docx"
    doc.save(output)
    print(f"Created: {output}")


if __name__ == "__main__":
    main()
